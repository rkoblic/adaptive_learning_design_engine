"""Output formatting utilities for display and download."""

import markdown
import re
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def markdown_to_html(md_content: str) -> str:
    """
    Convert markdown content to HTML for display.

    Args:
        md_content: Markdown string

    Returns:
        HTML string with proper formatting
    """
    # Configure markdown extensions for tables and code blocks
    extensions = [
        'markdown.extensions.tables',
        'markdown.extensions.fenced_code',
        'markdown.extensions.toc',
        'markdown.extensions.nl2br'
    ]

    md = markdown.Markdown(extensions=extensions)
    html = md.convert(md_content)

    return html


def prepare_download(data: dict) -> str:
    """
    Prepare filename for curriculum download.

    Args:
        data: Confirmed data dict with project and learner info

    Returns:
        Sanitized filename string
    """
    # Build filename from project and learner data
    project_title = data.get('project', {}).get('project_title', 'course')
    learner_name = data.get('learner', {}).get('learner_name', '')

    # Sanitize for filename
    project_title = sanitize_filename(project_title)
    learner_name = sanitize_filename(learner_name)

    if learner_name:
        filename = f"{project_title}_{learner_name}_curriculum.md"
    else:
        filename = f"{project_title}_curriculum.md"

    return filename


def sanitize_filename(text: str) -> str:
    """
    Sanitize text for use in filename.

    Args:
        text: Raw text string

    Returns:
        Sanitized string safe for filenames
    """
    # Remove or replace invalid characters
    text = re.sub(r'[<>:"/\\|?*]', '', text)
    # Replace spaces and multiple underscores
    text = re.sub(r'\s+', '_', text)
    text = re.sub(r'_+', '_', text)
    # Limit length
    text = text[:50].strip('_')
    return text or 'course'


def extract_toc_from_curriculum(md_content: str) -> list:
    """
    Extract table of contents from curriculum markdown.

    Args:
        md_content: Markdown string

    Returns:
        List of dicts with 'level', 'title', 'anchor' keys
    """
    toc = []
    lines = md_content.split('\n')

    for line in lines:
        # Match markdown headers
        match = re.match(r'^(#{1,3})\s+(.+)$', line)
        if match:
            level = len(match.group(1))
            title = match.group(2).strip()
            # Create anchor from title
            anchor = re.sub(r'[^a-z0-9\s-]', '', title.lower())
            anchor = re.sub(r'\s+', '-', anchor)

            toc.append({
                'level': level,
                'title': title,
                'anchor': anchor
            })

    return toc


def markdown_to_docx(md_content: str) -> io.BytesIO:
    """
    Convert markdown content to a Word document.

    Args:
        md_content: Markdown string

    Returns:
        BytesIO buffer containing the Word document
    """
    doc = Document()

    # Process markdown line by line
    lines = md_content.split('\n')
    current_list = None

    for line in lines:
        line = line.rstrip()

        # Skip empty lines
        if not line:
            current_list = None
            continue

        # Handle headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            current_list = None
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            current_list = None
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            current_list = None
        # Handle list items
        elif line.startswith('- '):
            text = line[2:]
            # Handle bold text in list items
            text = process_inline_formatting(text)
            p = doc.add_paragraph(text, style='List Bullet')
        elif line.startswith('* '):
            text = line[2:]
            text = process_inline_formatting(text)
            p = doc.add_paragraph(text, style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            text = process_inline_formatting(text)
            p = doc.add_paragraph(text, style='List Number')
        # Handle bold lines (like **Credits:** 3)
        elif line.startswith('**') and ':**' in line:
            p = doc.add_paragraph()
            # Extract the bold part and the value
            match = re.match(r'\*\*(.+?):\*\*\s*(.*)', line)
            if match:
                run = p.add_run(match.group(1) + ': ')
                run.bold = True
                p.add_run(match.group(2))
            else:
                p.add_run(process_inline_formatting(line))
        # Regular paragraph
        else:
            text = process_inline_formatting(line)
            p = doc.add_paragraph(text)

    # Save to BytesIO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


def process_inline_formatting(text: str) -> str:
    """
    Process inline markdown formatting (bold, italic) for plain text output.
    Note: For docx, we just strip the markers since python-docx handles runs differently.

    Args:
        text: Text with markdown formatting

    Returns:
        Clean text without markdown markers
    """
    # Remove bold markers
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Remove italic markers
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    return text
