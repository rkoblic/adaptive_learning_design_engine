"""File parsing utilities for extracting text from resume files."""

import io
from pypdf import PdfReader
from docx import Document as DocxDocument


def extract_text_from_file(file_storage) -> str:
    """
    Extract text from uploaded PDF or DOCX file.

    Args:
        file_storage: Flask FileStorage object from request.files

    Returns:
        Extracted text content as string

    Raises:
        ValueError: If file type is not supported
    """
    filename = file_storage.filename.lower()

    if filename.endswith('.pdf'):
        return extract_from_pdf(file_storage)
    elif filename.endswith('.docx'):
        return extract_from_docx(file_storage)
    else:
        raise ValueError(f"Unsupported file type: {filename}. Please upload a PDF or DOCX file.")


def extract_from_pdf(file_storage) -> str:
    """
    Extract text from PDF file.

    Args:
        file_storage: Flask FileStorage object

    Returns:
        Extracted text content
    """
    try:
        # Read file into bytes buffer
        file_bytes = io.BytesIO(file_storage.read())
        pdf_reader = PdfReader(file_bytes)

        text_parts = []
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Error reading PDF file: {str(e)}")


def extract_from_docx(file_storage) -> str:
    """
    Extract text from DOCX file.

    Args:
        file_storage: Flask FileStorage object

    Returns:
        Extracted text content
    """
    try:
        # Read file into bytes buffer
        file_bytes = io.BytesIO(file_storage.read())
        doc = DocxDocument(file_bytes)

        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Error reading DOCX file: {str(e)}")
